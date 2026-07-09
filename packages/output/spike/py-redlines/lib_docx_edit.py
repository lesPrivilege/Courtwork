"""按 instructions.json 对 python-docx Document 做文本层面的修改，
产出 Python-Redlines 需要的“修改后”版本。这一层完全在 python-docx 的
高层对象模型内完成（不需要碰 XML），因为增删改的对象都是脚本自建、
每段落只有一个 run 的简单文档。真实卷宗的段落经常被 Word 拆成多个
run（字体/修订历史造成的碎片化），这是留给正式管线的已知复杂度，不
在本 spike 范围内验证。
"""

import copy


def _iter_body_paragraphs(doc):
    return list(doc.paragraphs)


def _find_paragraph_by_quote(doc, quote):
    for p in _iter_body_paragraphs(doc):
        if quote in p.text:
            return p
    return None


def _find_table_cell(doc, row_contains, column_header=None):
    for table in doc.tables:
        header_cells = [c.text.strip() for c in table.rows[0].cells]
        for row in table.rows[1:]:
            row_texts = [c.text for c in row.cells]
            if any(row_contains in t for t in row_texts):
                if column_header is None:
                    return row, None
                col_idx = header_cells.index(column_header)
                return row, row.cells[col_idx]
    return None, None


def _find_table_row(doc, row_contains):
    for table in doc.tables:
        for row in table.rows[1:]:
            if any(row_contains in c.text for c in row.cells):
                return table, row
    return None, None


def _set_paragraph_text(paragraph, text):
    for run in list(paragraph.runs):
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = text
    else:
        paragraph.add_run(text)


def apply_instructions(doc, instructions):
    """就地修改 doc，返回 {instruction_id: {"status": ..., "quote_for_comment": ...}}"""
    results = {}

    for ins in instructions:
        iid = ins["id"]
        kind = ins["kind"]
        locator = ins["locator"]

        try:
            if kind == "replace" and "tableCell" in locator:
                row_contains = locator["tableCell"]["rowContains"]
                column_header = locator["tableCell"]["columnHeader"]
                row, cell = _find_table_cell(doc, row_contains, column_header)
                if cell is None:
                    results[iid] = {"status": "locator_not_found"}
                    continue
                para = cell.paragraphs[0]
                if locator["quote"] not in para.text:
                    results[iid] = {"status": "locator_text_mismatch"}
                    continue
                _set_paragraph_text(para, ins["replacementText"])
                results[iid] = {"status": "applied", "quote_for_comment": ins["replacementText"]}

            elif kind == "replace":
                para = _find_paragraph_by_quote(doc, locator["quote"])
                if para is None:
                    results[iid] = {"status": "locator_not_found"}
                    continue
                new_text = para.text.replace(locator["quote"], ins["replacementText"])
                _set_paragraph_text(para, new_text)
                results[iid] = {"status": "applied", "quote_for_comment": ins["replacementText"]}

            elif kind == "insert":
                anchor_text = locator["afterParagraphContaining"]
                anchor_para = _find_paragraph_by_quote(doc, anchor_text)
                if anchor_para is None:
                    results[iid] = {"status": "locator_not_found"}
                    continue
                new_p = copy.deepcopy(anchor_para._p)
                for r in new_p.findall(
                    ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t"
                ):
                    r.text = ""
                anchor_para._p.addnext(new_p)
                from docx.text.paragraph import Paragraph

                inserted = Paragraph(new_p, anchor_para._parent)
                lines = ins["insertText"].split("\n")
                inserted.add_run(lines[0]).bold = True
                for line in lines[1:]:
                    inserted.add_run("\n" + line)
                results[iid] = {"status": "applied", "quote_for_comment": lines[0]}

            elif kind == "delete" and "tableRowContains" in locator:
                table, row = _find_table_row(doc, locator["tableRowContains"])
                if row is None:
                    results[iid] = {"status": "locator_not_found"}
                    continue
                table._tbl.remove(row._tr)
                results[iid] = {"status": "applied", "quote_for_comment": None}

            elif kind == "delete":
                para = _find_paragraph_by_quote(doc, locator["quote"])
                if para is None:
                    results[iid] = {"status": "locator_not_found"}
                    continue
                para._element.getparent().remove(para._element)
                results[iid] = {"status": "applied", "quote_for_comment": None}

            elif kind == "comment-only":
                para = _find_paragraph_by_quote(doc, locator["quote"])
                if para is None:
                    results[iid] = {"status": "locator_not_found"}
                    continue
                results[iid] = {"status": "applied", "quote_for_comment": locator["quote"]}

            else:
                results[iid] = {"status": "unknown_kind"}

        except Exception as e:  # noqa: BLE001 - spike 阶段先把异常打平记录
            results[iid] = {"status": f"error: {e!r}"}

    return results
