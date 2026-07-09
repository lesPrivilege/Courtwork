"""生成 spike 用样例合同（虚构公司名/金额，脱敏）。"""

from docx import Document
from docx.shared import Pt

OUT_PATH = "original.docx"


def add_heading(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16)
    p.alignment = 1  # center


def add_clause(doc, title, body):
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    doc.add_paragraph(body)


def build():
    doc = Document()
    add_heading(doc, "买卖合同")

    doc.add_paragraph("甲方（买受人）：星辰科技有限公司")
    doc.add_paragraph("乙方（出卖人）：恒源贸易有限公司")

    add_clause(
        doc,
        "第一条 标的物",
        "乙方向甲方出售办公设备一批，规格、数量、单价详见本合同附表一。",
    )
    add_clause(
        doc,
        "第二条 价款",
        "本合同总价款为人民币叁拾万元整（￥300,000.00）。",
    )
    add_clause(
        doc,
        "第三条 交付期限",
        "乙方应于本合同签订之日起三十日内将标的物交付至甲方指定地点。",
    )

    p = doc.add_paragraph()
    p.add_run("第四条 付款方式").bold = True
    doc.add_paragraph("甲方应按照下表约定的进度支付价款：")

    table = doc.add_table(rows=4, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "期数"
    hdr[1].text = "支付比例"
    hdr[2].text = "支付时间"

    rows_data = [
        ("第一期", "30%", "合同签订之日起3日内"),
        ("第二期", "40%", "标的物交付之日起5日内"),
        ("第三期", "30%", "验收合格之日起10日内"),
    ]
    for i, (a, b, c) in enumerate(rows_data, start=1):
        cells = table.rows[i].cells
        cells[0].text = a
        cells[1].text = b
        cells[2].text = c

    doc.add_paragraph("")

    add_clause(
        doc,
        "第五条 质量保证",
        "乙方保证标的物符合国家相关质量标准，质保期为交付之日起壹年。",
    )
    add_clause(
        doc,
        "第六条 违约责任",
        "任何一方违反本合同约定，应向守约方支付合同总价款百分之十的违约金。",
    )
    add_clause(
        doc,
        "第七条 争议解决",
        "因本合同引起的争议，双方应协商解决；协商不成的，提交甲方所在地人民法院诉讼解决。",
    )
    add_clause(
        doc,
        "第八条 其他约定",
        "本合同一式两份，甲乙双方各执一份，自双方签字盖章之日起生效。",
    )

    doc.add_paragraph("")
    doc.add_paragraph("甲方（盖章）：__________          乙方（盖章）：__________")
    doc.add_paragraph("签订日期：__________              签订日期：__________")

    doc.save(OUT_PATH)
    print(f"wrote {OUT_PATH}")


if __name__ == "__main__":
    build()
